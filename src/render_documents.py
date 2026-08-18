import argparse
import json
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

TEMPLATE_RESUME = "assets/templates/resume_template.docx"
TEMPLATE_COVER = "assets/templates/cover_letter_template.docx"
BASE_RESUME = "assets/base_resume.json"
OUTPUT_DIR = "output"

COVER_FONT = "Aptos"
COVER_SIZE = 11


def _clear_body(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _add_hyperlink(paragraph, text: str, url: str, font_name: str = "Aptos", font_size: int = 11) -> None:
    """Add a styled hyperlink (blue, underlined) to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Apply Hyperlink character style (blue + underline)
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    # Font
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rPr.append(rFonts)

    # Size (half-points)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(font_size * 2))
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(font_size * 2))
    rPr.append(szCs)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _run(paragraph, text: str, bold: bool = False, size: float = 10, font: str | None = None):
    r = paragraph.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if font:
        r.font.name = font
    return r


def _spacing(para, before: float = 0, after: float = 0, line: float = 1.1) -> None:
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line



def _no_em_dash(text: str) -> str:
    return re.sub(r"\s*\u2014\s*", ", ", text)


# ── Resume constants ──────────────────────────────────────────────────────────
RF = "Lato"
HEADER_COLOR = RGBColor(0x36, 0x5F, 0x91)
CONTENT_PT = 10
HEADER_PT = 12
RIGHT_TAB = 10800  # 7.5" text width in twips (8.5" - 0.5" - 0.5" margins)


def _r_section(doc: Document, text: str):
    """Section header: Lato 12pt #365F91 bold + bottom border line."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = RF
    r.font.size = Pt(HEADER_PT)
    r.font.color.rgb = HEADER_COLOR
    # Bottom border (horizontal rule)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    btm = OxmlElement("w:bottom")
    btm.set(qn("w:val"), "single")
    btm.set(qn("w:sz"), "6")
    btm.set(qn("w:space"), "0")
    btm.set(qn("w:color"), "365F91")
    pBdr.append(btm)
    pPr.append(pBdr)
    _spacing(p, before=14, after=2, line=1.0)
    return p


def _r_right_tab(p):
    """Add a right-aligned tab stop at the right margin."""
    pPr = p._p.get_or_add_pPr()
    tabs_el = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(RIGHT_TAB))
    tabs_el.append(tab)
    pPr.append(tabs_el)


def _r_bullet(doc: Document, bold_part: str, normal_part: str = ""):
    """Bullet paragraph starting at the left margin with space after."""
    from docx.shared import Cm
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    # Bullet at left edge, text wraps slightly indented past bullet
    pf.left_indent = Cm(0.35)
    pf.first_line_indent = Cm(-0.35)
    rb_bull = p.add_run("\u2022  ")
    rb_bull.font.name = RF
    rb_bull.font.size = Pt(CONTENT_PT)
    if bold_part:
        rb = p.add_run(bold_part)
        rb.bold = True
        rb.font.name = RF
        rb.font.size = Pt(CONTENT_PT)
    if normal_part:
        rn = p.add_run(normal_part)
        rn.font.name = RF
        rn.font.size = Pt(CONTENT_PT)
    _spacing(p, before=0, after=4, line=1.0)
    return p


def build_resume(resume_json: dict, base_resume: dict, output_path: str) -> None:
    doc = Document(TEMPLATE_RESUME)
    _clear_body(doc)

    identity = base_resume["identity"]
    fixed = base_resume["fixed_sections"]

    # ── Name ──────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    r = p.add_run(identity["name"])
    r.bold = True
    r.font.name = RF
    r.font.size = Pt(18)
    r.font.color.rgb = HEADER_COLOR
    _spacing(p, after=2, line=1.0)

    # ── Career tags (first 3 from target_role_tags) ───────────────────────────
    tags = resume_json.get("target_role_tags", [])[:3]
    if tags:
        p = doc.add_paragraph()
        r = p.add_run("  |  ".join(tags))
        r.bold = True
        r.font.name = RF
        r.font.size = Pt(CONTENT_PT)
        _spacing(p, after=4, line=1.0)

    # ── Contact line: email(hyperlink) | phone | location | LinkedIn ──────────
    p = doc.add_paragraph()
    _add_hyperlink(p, identity["email"], "mailto:" + identity["email"], RF, CONTENT_PT)
    r2 = p.add_run("  |  M: " + identity["mobile"] + "  |  " + identity["location"] + "  |  ")
    r2.font.name = RF
    r2.font.size = Pt(CONTENT_PT)
    _add_hyperlink(p, "LinkedIn", identity["linkedin"], RF, CONTENT_PT)
    _spacing(p, after=6, line=1.0)

    # ── CAREER PROFILE ────────────────────────────────────────────────────────
    _r_section(doc, "CAREER PROFILE")
    p = doc.add_paragraph()
    r = p.add_run(_no_em_dash(resume_json["career_profile"]))
    r.font.name = RF
    r.font.size = Pt(CONTENT_PT)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _spacing(p, before=2, after=4, line=1.0)

    # ── AREAS OF EXPERTISE/HIGHLIGHTS ─────────────────────────────────────────
    _r_section(doc, "AREAS OF EXPERTISE/HIGHLIGHTS")
    for item in resume_json.get("areas_of_expertise", []):
        if isinstance(item, dict):
            bold_part = item.get("area", "")
            normal_part = ": " + item.get("description", "") if item.get("description") else ""
        else:
            # Backward compat: plain string
            if ":" in item:
                key, rest = item.split(":", 1)
                bold_part, normal_part = key + ":", rest
            else:
                bold_part, normal_part = item, ""
        _r_bullet(doc, _no_em_dash(bold_part), _no_em_dash(normal_part))

    # ── PROFESSIONAL EXPERIENCE ───────────────────────────────────────────────
    _r_section(doc, "PROFESSIONAL EXPERIENCE")
    for job in resume_json.get("professional_experience", []):
        # Title line with right-aligned date
        p = doc.add_paragraph()
        title_str = job["title"] + "  \u2015  " + job["company"]
        r = p.add_run(title_str)
        r.bold = True
        r.font.name = RF
        r.font.size = Pt(11)
        period = job.get("period", "")
        if period:
            _r_right_tab(p)
            r2 = p.add_run("\t" + period)
            r2.bold = True
            r2.font.name = RF
            r2.font.size = Pt(11)
        _spacing(p, before=10, after=2, line=1.0)

        # Summary
        if job.get("tailored_summary"):
            p = doc.add_paragraph()
            r = p.add_run(_no_em_dash(job["tailored_summary"]))
            r.font.name = RF
            r.font.size = Pt(CONTENT_PT)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _spacing(p, after=6, line=1.0)

        # Key Achievements
        achievements = job.get("tailored_achievements", [])
        if achievements:
            p = doc.add_paragraph()
            r = p.add_run("Key Achievements")
            r.bold = True
            r.italic = True
            r.font.name = RF
            r.font.size = Pt(CONTENT_PT)
            _spacing(p, before=2, after=1, line=1.0)

            for ach in achievements:
                if isinstance(ach, dict):
                    bold_part = ach.get("highlight", "")
                    normal_part = " " + ach.get("detail", "") if ach.get("detail") else ""
                else:
                    # Backward compat: plain string — split on first " by/through/via"
                    m = re.search(r"( by | through | via | resulting | enabling )", ach, re.I)
                    if m and m.start() > 8:
                        bold_part = ach[:m.start()]
                        normal_part = ach[m.start():]
                    else:
                        words = ach.split()
                        bold_part = " ".join(words[:6])
                        normal_part = (" " + " ".join(words[6:])) if len(words) > 6 else ""
                _r_bullet(doc, _no_em_dash(bold_part), _no_em_dash(normal_part))

    # ── EDUCATION ─────────────────────────────────────────────────────────────
    _r_section(doc, "EDUCATION")
    for edu in fixed.get("education", []):
        p = doc.add_paragraph()
        r = p.add_run(edu["qualification"])
        r.bold = True
        r.font.name = RF
        r.font.size = Pt(CONTENT_PT)
        inst = p.add_run("  \u2013  " + edu["institution"])
        inst.font.name = RF
        inst.font.size = Pt(CONTENT_PT)
        if edu.get("period"):
            _r_right_tab(p)
            rt = p.add_run("\t" + edu["period"])
            rt.font.name = RF
            rt.font.size = Pt(CONTENT_PT)
        _spacing(p, before=3, after=4, line=1.0)
        if edu.get("details"):
            p = doc.add_paragraph()
            p.add_run(edu["details"]).font.name = RF
            p.runs[-1].font.size = Pt(CONTENT_PT)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _spacing(p, after=4, line=1.0)

    for cert in fixed.get("certifications", []):
        # Extract trailing (YYYY) and place right-aligned
        m = re.search(r"\s*\((\d{4})\)\s*$", cert)
        cert_text = cert[:m.start()].strip() if m else cert
        cert_year = m.group(1) if m else ""
        p = doc.add_paragraph()
        r = p.add_run(cert_text)
        r.bold = True
        r.font.name = RF
        r.font.size = Pt(CONTENT_PT)
        if cert_year:
            _r_right_tab(p)
            ry = p.add_run("\t" + cert_year)
            ry.font.name = RF
            ry.font.size = Pt(CONTENT_PT)
        _spacing(p, after=4, line=1.0)

    # ── REFERENCES ────────────────────────────────────────────────────────────
    _r_section(doc, "REFERENCES")
    p = doc.add_paragraph()
    r = p.add_run(fixed.get("references", "References available upon request."))
    r.font.name = RF
    r.font.size = Pt(CONTENT_PT)
    _spacing(p, before=2, line=1.0)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Resume saved: {output_path}")


def build_cover_letter(cover_letter_json: dict, base_resume: dict, output_path: str) -> None:
    doc = Document(TEMPLATE_COVER)
    _clear_body(doc)

    identity = base_resume["identity"]
    today = date.today().strftime("%d %B %Y").lstrip("0")
    S = COVER_SIZE
    F = COVER_FONT

    def cp(text="", bold=False, justify=False, before=0, after=8, line=1.0):
        """Add a cover letter paragraph with standard formatting."""
        p = doc.add_paragraph()
        if text:
            r = p.add_run(_no_em_dash(text))
            r.bold = bold
            r.font.size = Pt(S)
            r.font.name = F
        if justify:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _spacing(p, before=before, after=after, line=line)
        return p

    # --- Header block ---
    p = doc.add_paragraph()
    # Name (bold)
    r = p.add_run(identity["name"])
    r.bold = True
    r.font.size = Pt(S)
    r.font.name = F
    # Location
    r2 = p.add_run("\n" + identity["location"])
    r2.font.size = Pt(S)
    r2.font.name = F
    # Email as mailto hyperlink + phone
    r3 = p.add_run("\n")
    r3.font.size = Pt(S)
    _add_hyperlink(p, identity["email"], "mailto:" + identity["email"], F, S)
    r4 = p.add_run(" | " + identity["mobile"])
    r4.font.size = Pt(S)
    r4.font.name = F
    # LinkedIn on its own line
    r5 = p.add_run("\n")
    r5.font.size = Pt(S)
    _add_hyperlink(p, "LinkedIn", identity["linkedin"], F, S)
    _spacing(p, before=0, after=12, line=1.0)

    # --- Blank line after header ---
    cp()

    # --- Date ---
    cp(today)

    # --- Recipient: company name (not "Hiring Team") ---
    company = cover_letter_json.get("company_name", "")
    hiring_manager = cover_letter_json.get("hiring_manager_name", "")
    if hiring_manager and company:
        recipient = hiring_manager + "\n" + company
    elif company:
        recipient = company
    elif hiring_manager:
        recipient = hiring_manager
    else:
        recipient = ""
    if recipient:
        p = doc.add_paragraph()
        r = p.add_run(recipient)
        r.font.size = Pt(S)
        r.font.name = F
        _spacing(p, after=8, line=1.15)

    # --- Salutation ---
    cp(cover_letter_json["salutation"])

    # --- Space after salutation ---
    cp()

    # --- Body paragraphs (justified, with space after each) ---
    for field in ["opening_paragraph", "body_paragraph_1", "body_paragraph_2", "closing_paragraph"]:
        cp(_no_em_dash(cover_letter_json[field]), justify=True, after=10)

    # --- Space before sign-off ---
    cp()

    # --- Sign-off ---
    cp(cover_letter_json["sign_off"], after=0)

    # --- Name ---
    p = doc.add_paragraph()
    r = p.add_run(identity["name"])
    r.bold = True
    r.font.size = Pt(S)
    r.font.name = F
    _spacing(p)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Cover letter saved: {output_path}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--resume-json", required=True, help="Path to resume JSON file")
    parser.add_argument("--cover-letter-json", required=True, help="Path to cover letter JSON file")
    parser.add_argument("--job-id", required=True, help="Job ID used for output filenames")
    args = parser.parse_args()

    with open(args.resume_json, encoding="utf-8") as f:
        resume_json = json.load(f)
    with open(args.cover_letter_json, encoding="utf-8") as f:
        cover_letter_json = json.load(f)
    with open(BASE_RESUME, encoding="utf-8") as f:
        base_resume = json.load(f)

    build_resume(resume_json, base_resume, f"{OUTPUT_DIR}/{args.job_id}_resume.docx")
    build_cover_letter(cover_letter_json, base_resume, f"{OUTPUT_DIR}/{args.job_id}_cover_letter.docx")


if __name__ == "__main__":
    main()
